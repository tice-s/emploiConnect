"""Extraction de texte brut à partir d'un CV importé (PDF ou DOCX)."""
import logging

logger = logging.getLogger("ai_engine")


class DocumentParsingError(Exception):
    pass


def extract_text_from_file(file_field) -> str:
    """
    Extrait le texte brut d'un fichier CV (PDF ou DOCX) déjà enregistré sur disque
    (FileField Django). Lève DocumentParsingError si le format n'est pas supporté
    ou si l'extraction échoue.
    """
    name = file_field.name.lower()
    file_field.open("rb")
    try:
        if name.endswith(".pdf"):
            return _extract_pdf(file_field)
        elif name.endswith(".docx"):
            return _extract_docx(file_field)
        raise DocumentParsingError("Format de fichier non supporté (PDF ou DOCX uniquement).")
    finally:
        file_field.close()


def _extract_pdf(file_field) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise DocumentParsingError("La librairie pypdf n'est pas installée.") from exc

    try:
        reader = PdfReader(file_field)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        logger.exception("Échec de l'extraction PDF")
        raise DocumentParsingError("Impossible de lire ce fichier PDF.") from exc

    if not text.strip():
        raise DocumentParsingError(
            "Aucun texte n'a pu être extrait de ce PDF (probablement un scan/image sans OCR)."
        )
    return text


def _extract_docx(file_field) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise DocumentParsingError("La librairie python-docx n'est pas installée.") from exc

    try:
        document = docx.Document(file_field)
        paragraphs = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                paragraphs.extend(cell.text for cell in row.cells)
        text = "\n".join(p for p in paragraphs if p.strip())
    except Exception as exc:
        logger.exception("Échec de l'extraction DOCX")
        raise DocumentParsingError("Impossible de lire ce fichier DOCX.") from exc

    if not text.strip():
        raise DocumentParsingError("Aucun texte n'a pu être extrait de ce document.")
    return text
